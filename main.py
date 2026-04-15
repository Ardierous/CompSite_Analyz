from flask import Flask, render_template, request, jsonify, send_file, Response
from flask_cors import CORS
import os
import threading
from datetime import datetime
from urllib.parse import urlparse, quote, urljoin, urlunparse
import warnings
import io
import requests
import json
from pathlib import Path
import sys
import platform
import subprocess
import socket
import atexit
import signal
import unicodedata
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
import time
import tempfile
warnings.filterwarnings('ignore')

_executor_md_docx = ThreadPoolExecutor(max_workers=2, thread_name_prefix='md2docx')
MD_DOCX_TIMEOUT = 120

# Путь к PID файлу
PID_FILE = Path(__file__).parent / '.app.pid'

# ANSI коды для цветного вывода в терминале
class Colors:
    """ANSI escape коды для цветного вывода"""
    GREEN = '\033[92m'  # Зеленый
    RED = '\033[91m'    # Красный
    YELLOW = '\033[93m' # Желтый
    BLUE = '\033[94m'   # Синий
    RESET = '\033[0m'   # Сброс цвета
    BOLD = '\033[1m'    # Жирный

def print_green(text):
    """Выводит текст зеленым цветом"""
    print(f"{Colors.GREEN}{text}{Colors.RESET}")

def print_red(text):
    """Выводит текст красным цветом"""
    print(f"{Colors.RED}{text}{Colors.RESET}")

def print_yellow(text):
    """Выводит текст желтым цветом"""
    print(f"{Colors.YELLOW}{text}{Colors.RESET}")

def print_blue(text):
    """Выводит текст синим цветом"""
    print(f"{Colors.BLUE}{text}{Colors.RESET}")

def is_running_in_docker():
    """Проверяет, запущено ли приложение в контейнере Docker"""
    # Проверяем наличие файла /.dockerenv (стандартный способ)
    if Path('/.dockerenv').exists():
        return True
    # Проверяем переменную окружения (если установлена)
    if os.getenv('DOCKER_CONTAINER') == 'true':
        return True
    # Проверяем cgroup (альтернативный способ)
    try:
        with open('/proc/self/cgroup', 'r') as f:
            if 'docker' in f.read():
                return True
    except:
        pass
    return False

def check_port_in_use(host, port):
    """Проверяет, занят ли порт"""
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.settimeout(1)
    result = sock.connect_ex((host, port))
    sock.close()
    return result == 0

def is_process_running(pid):
    """Проверяет, запущен ли процесс с указанным PID"""
    try:
        if platform.system() == 'Windows':
            result = subprocess.run(
                ['tasklist', '/FI', f'PID eq {pid}'],
                capture_output=True,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
            return str(pid) in result.stdout
        else:
            # Linux/Mac: используем kill -0 для проверки существования процесса
            result = subprocess.run(['kill', '-0', str(pid)], capture_output=True)
            return result.returncode == 0
    except:
        return False

def kill_process(pid):
    """Завершает процесс с указанным PID"""
    try:
        if platform.system() == 'Windows':
            subprocess.run(
                ['taskkill', '/F', '/PID', str(pid)],
                capture_output=True,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
        else:
            subprocess.run(['kill', '-9', str(pid)], capture_output=True)
        return True
    except Exception as e:
        logger.warning(f"Не удалось завершить процесс {pid}: {e}")
        return False

def kill_process_on_port(port, exclude_pid=None):
    """Завершает все процессы, использующие указанный порт
    
    Args:
        port: Порт для проверки
        exclude_pid: PID процесса, который нужно исключить из проверки (текущий процесс)
    """
    # В контейнере Docker не проверяем порт - Docker сам управляет портами
    if is_running_in_docker():
        logger.info("Запуск в контейнере Docker - пропускаем проверку порта")
        return False
    
    if exclude_pid is None:
        exclude_pid = os.getpid()  # Исключаем текущий процесс
    
    killed_count = 0
    try:
        if platform.system() == 'Windows':
            # Windows: используем netstat и taskkill
            # Сначала получаем все PID процессов на порту
            result = subprocess.run(
                ['netstat', '-ano'],
                capture_output=True,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
            pids_to_kill = set()
            for line in result.stdout.split('\n'):
                if f':{port}' in line and 'LISTENING' in line:
                    parts = line.split()
                    if len(parts) > 4:
                        pid = parts[-1]
                        if pid.isdigit():
                            pids_to_kill.add(pid)
            
            # Завершаем все найденные процессы
            for pid in pids_to_kill:
                if kill_process(pid):
                    logger.info(f"Завершен процесс с PID {pid}, занимавший порт {port}")
                    print_red(f"✓ Завершен процесс с PID {pid}, занимавший порт {port}")
                    killed_count += 1
        else:
            # Linux/Mac: используем ss или netstat (более универсально, чем lsof)
            # Сначала пробуем ss (более современная утилита)
            result = subprocess.run(
                ['ss', '-tlnp'],
                capture_output=True,
                text=True
            )
            if result.returncode == 0:
                # Парсим вывод ss
                import re
                for line in result.stdout.split('\n'):
                    if f':{port}' in line and 'LISTEN' in line:
                        # Извлекаем PID из строки (формат: pid=12345)
                        pid_match = re.search(r'pid=(\d+)', line)
                        if pid_match:
                            pid = pid_match.group(1)
                            # Исключаем текущий процесс и его дочерние процессы
                            if pid != str(exclude_pid) and kill_process(pid):
                                logger.info(f"Завершен процесс с PID {pid}, занимавший порт {port}")
                                print_red(f"✓ Завершен процесс с PID {pid}, занимавший порт {port}")
                                killed_count += 1
            else:
                # Если ss недоступен, пробуем netstat
                result = subprocess.run(
                    ['netstat', '-tlnp'],
                    capture_output=True,
                    text=True
                )
                if result.returncode == 0:
                    for line in result.stdout.split('\n'):
                        if f':{port}' in line and 'LISTEN' in line:
                            parts = line.split()
                            # PID обычно в последнем столбце, формат: 12345/python
                            if len(parts) > 0:
                                last_part = parts[-1]
                                pid = last_part.split('/')[0]
                                if pid.isdigit() and pid != str(exclude_pid):
                                    # Исключаем текущий процесс и его дочерние процессы
                                    if kill_process(pid):
                                        logger.info(f"Завершен процесс с PID {pid}, занимавший порт {port}")
                                        print_red(f"✓ Завершен процесс с PID {pid}, занимавший порт {port}")
                                        killed_count += 1
    except Exception as e:
        logger.error(f"Ошибка при попытке завершить процесс на порту {port}: {e}")
        print(f"⚠ Предупреждение: не удалось автоматически завершить процесс на порту {port}")
    
    return killed_count > 0

def check_and_kill_existing_instance():
    """Проверяет наличие запущенного экземпляра приложения и завершает его"""
    if not PID_FILE.exists():
        return True  # PID файл не существует, приложение не запущено
    
    try:
        # Читаем PID из файла
        with open(PID_FILE, 'r') as f:
            old_pid_str = f.read().strip()
            if not old_pid_str:
                # Пустой файл, удаляем его
                PID_FILE.unlink()
                return True
            old_pid = int(old_pid_str)
        
        # Проверяем, запущен ли процесс
        if is_process_running(old_pid):
            logger.warning(f"Обнаружен запущенный экземпляр приложения (PID: {old_pid}). Принудительно завершаю...")
            print_yellow(f"⚠ Обнаружен запущенный экземпляр приложения (PID: {old_pid}). Принудительно завершаю...")
            
            # Принудительно завершаем процесс
            if kill_process(old_pid):
                logger.info(f"Предыдущий экземпляр приложения (PID: {old_pid}) успешно завершен")
                print_red(f"✓ Предыдущий экземпляр приложения (PID: {old_pid}) успешно завершен")
                # Ждем немного, чтобы процесс завершился
                import time
                time.sleep(2)  # Увеличено время ожидания
                
                # Проверяем еще раз, что процесс действительно завершен
                if is_process_running(old_pid):
                    logger.warning(f"Процесс {old_pid} все еще запущен после попытки завершения. Повторная попытка...")
                    print_yellow(f"⚠ Процесс {old_pid} все еще запущен. Повторная попытка завершения...")
                    kill_process(old_pid)
                    time.sleep(1)
            else:
                logger.error(f"Не удалось завершить предыдущий экземпляр (PID: {old_pid})")
                print_red(f"❌ Не удалось завершить предыдущий экземпляр (PID: {old_pid})")
                return False
        else:
            # Процесс не запущен, но PID файл остался (возможно, аварийное завершение)
            logger.info(f"PID файл найден, но процесс {old_pid} не запущен. Удаляю старый PID файл.")
            print(f"ℹ Обнаружен старый PID файл (процесс не запущен). Очищаю...")
        
        # Удаляем старый PID файл
        try:
            PID_FILE.unlink()
        except:
            pass
        
        return True
    except ValueError:
        # Неверный формат PID файла
        logger.warning("PID файл содержит неверные данные. Удаляю...")
        try:
            PID_FILE.unlink()
        except:
            pass
        return True
    except Exception as e:
        logger.error(f"Ошибка при проверке существующего экземпляра: {e}")
        # Пытаемся удалить поврежденный PID файл
        try:
            PID_FILE.unlink()
        except:
            pass
        return True  # Продолжаем запуск, даже если не удалось проверить

def create_pid_file():
    """Создает PID файл с текущим PID процесса"""
    try:
        current_pid = os.getpid()
        with open(PID_FILE, 'w') as f:
            f.write(str(current_pid))
        logger.info(f"Создан PID файл: {PID_FILE} (PID: {current_pid})")
    except Exception as e:
        logger.error(f"Не удалось создать PID файл: {e}")

def remove_pid_file():
    """Удаляет PID файл при завершении приложения"""
    try:
        if PID_FILE.exists():
            PID_FILE.unlink()
            logger.info("PID файл удален")
    except Exception as e:
        logger.error(f"Не удалось удалить PID файл: {e}")

# ВАЖНО: Отключаем телеметрию CrewAI ДО загрузки переменных окружения
# Это должно быть сделано как можно раньше, до импорта Agents_crew
os.environ["CREWAI_TELEMETRY_OPT_OUT"] = "1"
os.environ["CREWAI_TRACING_ENABLED"] = "false"

# Загрузка переменных окружения из .env файла
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # Если python-dotenv не установлен, используем только переменные окружения системы
    pass

# Убеждаемся, что телеметрия отключена (на случай, если .env переопределил значения)
os.environ["CREWAI_TELEMETRY_OPT_OUT"] = "1"
os.environ["CREWAI_TRACING_ENABLED"] = "false"

# Предзагрузка CrewAI в главном потоке — иначе при первом анализе (в worker-потоке)
# телеметрия CrewAI пытается зарегистрировать signal.signal() и падает с
# ValueError: signal only works in main thread
try:
    from Agents_crew import crew  # noqa: F401
except Exception:
    pass

# Вспомогательная функция для записи в debug.log (опционально)
def update_progress_safely(task_id, new_progress, message=None):
    """
    Безопасно обновляет прогресс задачи, гарантируя, что он только увеличивается
    
    Args:
        task_id: ID задачи
        new_progress: Новое значение прогресса (0-100)
        message: Опциональное сообщение для обновления
    """
    if task_id not in analysis_status:
        return
    
    current_prog = analysis_status[task_id].get('progress', 0)
    final_progress = max(current_prog, min(new_progress, 100))
    analysis_status[task_id]['progress'] = final_progress
    
    if message:
        analysis_status[task_id]['message'] = message

def write_debug_log(data):
    """Безопасно записывает данные в debug.log, если файл доступен"""
    try:
        debug_log_path = os.getenv('DEBUG_LOG_PATH')
        if not debug_log_path:
            # Пытаемся использовать путь по умолчанию, если он существует
            default_path = Path(__file__).parent / '.cursor' / 'debug.log'
            if default_path.parent.exists():
                debug_log_path = str(default_path)
            else:
                return  # Не записываем, если путь недоступен
        
        with open(debug_log_path, 'a', encoding='utf-8') as f:
            f.write(json.dumps(data) + '\n')
    except:
        pass  # Игнорируем ошибки записи в debug.log

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.text.run import Run
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    RELATIONSHIP_TYPE = None
    Run = None

import re


# Символы и эмодзи, которые часто используются как маркеры списка (буллиты)
_BULLET_CHARS = (
    '•', '●', '○', '▪', '▫', '◦', '∙', '►', '⬤', '➤', '→',
    '🟢', '🔵', '🔴', '🟡', '⬜', '🟦', '🟥', '⭐', '✅', '❌', '❗', '▪️',
)
_BULLET_CHARS_SORTED = tuple(sorted(_BULLET_CHARS, key=lambda s: -len(s)))


def _strip_leading_bullet_chars(text):
    """Убирает из начала текста все символы-буллиты (• ✅ и т.д.), чтобы в DOCX оставался только читаемый текст."""
    if not text or not text.strip():
        return text.strip()
    t = text.strip()
    while t:
        found = False
        for sym in _BULLET_CHARS_SORTED:
            if t.startswith(sym):
                t = t[len(sym):].lstrip()
                found = True
                break
        if not found:
            break
    return t


def _replace_md_images(text):
    """Заменяет синтаксис картинок ![alt](url) на текст (alt или [изображение]). Ограничения длины — против backtracking."""
    return re.sub(r'!\[([^\]]{0,500})\]\([^)]{1,2000}\)', lambda m: ('[' + m.group(1) + ']') if m.group(1).strip() else '[изображение]', text)


def _line_starts_with_emoji_bullet(stripped):
    """Проверяет, начинается ли строка с эмодзи/символа-буллета и продолжается текстом.
    Возвращает (True, bullet, rest) или (False, None, None); bullet — сам символ буллета для отображения в DOCX."""
    if not stripped or len(stripped) < 2:
        return False, None, None
    # Сравниваем с более длинными маркерами первыми (▪️ до ▪)
    for sym in _BULLET_CHARS_SORTED:
        if stripped.startswith(sym):
            rest = stripped[len(sym):].lstrip()
            if rest or sym in ('•', '●', '○', '▪', '▫'):
                return True, sym, rest
            return False, None, None
    # Один символ из категории "Symbol, other" (эмодзи, спецсимволы) + пробелы + текст
    import unicodedata
    first = stripped[0]
    if ord(first) > 0x1F and not first.isalnum() and first not in '#*-`':
        if unicodedata.category(first) == 'So' and len(stripped) > 1:
            rest = stripped[1:].lstrip()
            if rest:
                return True, first, rest
    return False, None, None


def _is_emoji_char(c):
    """Проверка, что символ — эмодзи/декоративный (удаляем из вывода)."""
    if len(c) != 1:
        return False
    if c in _BULLET_CHARS:
        return True
    o = ord(c)
    if 0x1F300 <= o <= 0x1F9FF or 0x2600 <= o <= 0x26FF or 0x2700 <= o <= 0x27BF:
        return True
    if 0x2300 <= o <= 0x23FF or 0x2B00 <= o <= 0x2BFF:
        return True
    if 0x25A0 <= o <= 0x25FF:
        return True
    if 0x2022 <= o <= 0x2023:
        return True
    if o in (0xFE0F, 0xFE0E):
        return True
    if o >= 0x1F000:
        return True
    return unicodedata.category(c) == 'So'


def _strip_emoji(text):
    """Удаляет эмодзи и символы категории So из текста. Убирает только пробел, оставшийся после удалённого эмодзи (не трогает пробелы между словами)."""
    if not text:
        return text
    result = []
    i = 0
    while i < len(text):
        c = text[i]
        if _is_emoji_char(c):
            i += 1
            if i < len(text) and text[i] == ' ':
                i += 1
            continue
        result.append(c)
        i += 1
    return ''.join(result)


def _add_run_with_emoji_font(paragraph, text, bold=False, italic=False, strike=False, font_name=None):
    """Добавляет текст в параграф. Эмодзи удаляются. Длинное тире заменяется на обычное."""
    if not text:
        return
    text = text.replace('\u2014', '\u2013')  # — (em dash) → – (en dash)
    text = _strip_emoji(text)
    if not text:
        return
    base_font = font_name or 'Segoe UI'
    run = paragraph.add_run(text)
    run.font.name = base_font
    run.bold = bold
    run.italic = italic
    run.font.strike = strike


def _normalize_url(url):
    """Добавляет https:// для URL без схемы (www.example.com → https://www.example.com)."""
    if not url or not url.strip():
        return url
    u = url.strip()
    if not u.startswith(('http://', 'https://', 'mailto:', '#')):
        if u.startswith('www.'):
            return 'https://' + u
        if u.startswith(('/')):
            return u  # относительный путь — не нормализуем
    return u


def _extract_urls_from_markdown(text):
    """Извлекает все URL из markdown: [text](url) и голые https?://... и www...."""
    urls = set()
    if not text:
        return urls
    def _add(u):
        u = _normalize_url(u)
        if u and not u.startswith('#') and not u.startswith('mailto:'):
            urls.add(u)
            if u.endswith('/'):
                urls.add(u.rstrip('/'))
            else:
                urls.add(u + '/')
    for m in _RE_LINK.finditer(text):
        _add(m.group(2).strip())
    for m in _RE_BARE_URL.finditer(text):
        _add(m.group(1))
    return urls


def _hostname_base(host):
    if not host:
        return ''
    h = host.lower()
    if h.startswith('www.'):
        return h[4:]
    return h


def _canonical_url_key(url):
    """Ключ URL для whitelist: хост без www, путь без лишнего слэша, query."""
    u = _normalize_url((url or '').strip())
    p = urlparse(u)
    host = _hostname_base(p.hostname or '')
    path = (p.path or '/').rstrip('/') or '/'
    q = p.query or ''
    return (host, path, q)


def _expand_url_variants_for_verified_set(u):
    """Варианты строки URL для множества verified_urls (www, слэш, http/https)."""
    u = _normalize_url((u or '').strip())
    out = {u} if u else set()
    if not u.startswith(('http://', 'https://')):
        return out
    p = urlparse(u)
    path = p.path or '/'
    if u.endswith('/') and path not in ('', '/'):
        out.add(u.rstrip('/'))
    elif path not in ('', '/') and not u.endswith('/'):
        out.add(u + '/')
    net = p.netloc.lower()
    if net.startswith('www.'):
        alt = urlunparse((p.scheme, net[4:], p.path, '', p.query, ''))
    else:
        alt = urlunparse((p.scheme, 'www.' + net, p.path, '', p.query, ''))
    out.add(_normalize_url(alt))
    for s in tuple(out):
        if s.startswith('https://'):
            out.add('http://' + s[8:])
        elif s.startswith('http://'):
            out.add('https://' + s[7:])
    return {x for x in out if x}


def _discover_internal_site_urls(company_url, max_pages=24, timeout=12):
    """
    BFS по ссылкам того же домена: реальные пути из HTML (как в меню сайта).
    Возвращает (множество ключей _canonical_url_key, множество строк для verified_urls).
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return set(), set()

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
        'Connection': 'keep-alive',
    }
    seed = _normalize_url(company_url)
    base_host = _hostname_base((urlparse(seed).hostname or ''))
    if not base_host:
        return set(), set()

    allowed_keys = set()
    verified_strings = set()

    def register_url(u):
        u = _normalize_url(u)
        allowed_keys.add(_canonical_url_key(u))
        verified_strings.update(_expand_url_variants_for_verified_set(u))

    queue = [seed]
    queued_keys = {_canonical_url_key(seed)}
    pages_fetched = 0

    while queue and pages_fetched < max_pages:
        url = queue.pop(0)
        try:
            r = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
        except requests.RequestException:
            continue
        pages_fetched += 1
        if r.status_code >= 400:
            continue
        final = r.url
        register_url(final)

        ct = (r.headers.get('Content-Type') or '').lower()
        if 'html' not in ct and not ct.startswith('text/') and 'application/xhtml' not in ct:
            continue
        try:
            soup = BeautifulSoup(r.text, 'html.parser')
        except Exception:
            continue
        for a in soup.find_all('a', href=True):
            href = (a.get('href') or '').strip()
            if not href or href.startswith(('#', 'javascript:', 'mailto:', 'tel:', 'data:')):
                continue
            abs_u = urljoin(final, href)
            p = urlparse(abs_u)
            if p.scheme not in ('http', 'https'):
                continue
            if _hostname_base(p.hostname or '') != base_host:
                continue
            clean = urlunparse((p.scheme, p.netloc, p.path, '', p.query, ''))
            clean = _normalize_url(clean)
            register_url(clean)
            lk = _canonical_url_key(clean)
            if lk not in queued_keys and len(queued_keys) < max_pages * 6:
                queued_keys.add(lk)
                queue.append(clean)

    return allowed_keys, verified_strings


def _sanitize_markdown_links(text, allowed_keys):
    """Убирает markdown-ссылки [текст](url), если url нет в whitelist реальных путей."""
    if not text or not allowed_keys:
        return text

    def repl(m):
        label, url = m.group(1), m.group(2).strip()
        if _canonical_url_key(url) in allowed_keys:
            return m.group(0)
        return label

    return re.sub(r'\[([^\]]*)\]\(([^)]+)\)', repl, text)


def _url_is_verified(url, verified_urls, company_url=None):
    """Можно ли сделать ссылку кликабельной: только проверенные или того же домена."""
    if not url or not url.strip():
        return False
    url_norm = _normalize_url(url.strip())
    if url_norm.startswith('/') and company_url:
        try:
            cmp = urlparse(company_url)
            base = f"{cmp.scheme or 'https'}://{cmp.netloc}"
            url_norm = base.rstrip('/') + ('/' if not url_norm.startswith('/') else '') + url_norm
        except Exception:
            pass
    if verified_urls is not None:
        return url_norm in verified_urls
    if company_url:
        try:
            cmp = urlparse(company_url)
            lnk = urlparse(url_norm)
            if lnk.netloc and cmp.netloc:
                cmp_domain = cmp.netloc.lower().replace('www.', '')
                lnk_domain = lnk.netloc.lower().replace('www.', '')
                return lnk_domain == cmp_domain or lnk_domain.endswith('.' + cmp_domain)
        except Exception:
            pass
    return False


def _add_hyperlink(paragraph, text, url):
    """Добавляет кликабельную гиперссылку. Только текст (без URL), чёрный цвет + подчёркивание — видно при печати."""
    if not DOCX_AVAILABLE or not text or not url or Run is None:
        _add_run_with_emoji_font(paragraph, text or '')
        return
    try:
        url = _normalize_url(url)
        text = text.replace('\u2014', '\u2013')  # длинное тире → обычное
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        r_elem = OxmlElement('w:r')
        new_run = Run(r_elem, paragraph)
        new_run.text = text
        new_run.font.name = 'Segoe UI'
        new_run.font.color.rgb = RGBColor(0, 0, 0)
        new_run.font.underline = True
        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)
    except Exception:
        _add_run_with_emoji_font(paragraph, text)


# ASCII-символы, экранируемые обратным слэшем в CommonMark
_ESCAPABLE = set(r'!"#$%&\'()*+,\-./:;<=>?@[\]^_`{|}~')
# Предкомпилированный regex для инлайн-разметки (вызывается на каждой строке)
_RE_INLINE = re.compile(
    r'(\*\*\*[^*]{1,300}\*\*\*|___[^_]{1,300}___|\*\*[^*]{1,300}\*\*|__[^_]{1,300}__|'
    r'\*[^*]{1,300}\*|_[^_]{1,300}_|`[^`]{1,200}`|~~[^~]{1,200}~~|'
    r'\[\^[^\]]+\]|\[[^\]]{1,200}\]\([^)]{1,1500}\)|[^*`~\[]+)'
)
_RE_LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
# Голые URL (http/https/www) для кликабельных ссылок
_RE_BARE_URL = re.compile(r'(https?://[^\s\)\]<>"\']+|www\.[^\s\)\]<>"\']+)')
_RE_FOOTNOTE_REF = re.compile(r'\[\^([^\]]+)\]')
_RE_FOOTNOTE_DEF = re.compile(r'^\[\^([^\]]+)\]:\s*(.*)$')


def _text_for_width_calc(text):
    """Текст для расчёта ширины: [текст](url) → текст (отображаемая длина ссылки)."""
    if not text:
        return text
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def _apply_backslash_escapes(text):
    """Обрабатывает экранирование \\ и \\X по CommonMark: \\ → \, \\* → * и т.д."""
    result = []
    i = 0
    while i < len(text):
        if text[i] == '\\' and i + 1 < len(text):
            nxt = text[i + 1]
            if nxt in _ESCAPABLE:
                result.append(nxt)
                i += 2
                continue
            if nxt == '\\':
                result.append('\\')
                i += 2
                continue
        result.append(text[i])
        i += 1
    return ''.join(result)


def _split_and_add_with_bare_urls(paragraph, text, footnote_ctx=None):
    """Разбивает текст по голым URL и добавляет в параграф; URL делаются кликабельными только если проверены."""
    if not text or not _RE_BARE_URL.search(text):
        _add_run_with_emoji_font(paragraph, text)
        return
    ctx = footnote_ctx or {}
    verified = ctx.get('verified_urls')
    company_url = ctx.get('company_url')
    last = 0
    for m in _RE_BARE_URL.finditer(text):
        if m.start() > last:
            _add_run_with_emoji_font(paragraph, text[last:m.start()])
        url = _normalize_url(m.group(1))
        if _url_is_verified(url, verified, company_url):
            _add_hyperlink(paragraph, m.group(1), url)
        else:
            _add_run_with_emoji_font(paragraph, m.group(1))
        last = m.end()
    if last < len(text):
        _add_run_with_emoji_font(paragraph, text[last:])


def _add_inline_formatted(paragraph, text, footnote_ctx=None):
    """Добавляет в параграф текст с поддержкой **/__ жирный, */_ курсив, `код`, ~~зачёркнутый~~, [текст](url), [^id] сноски; экранирование \\; картинки → подпись."""
    text = _replace_md_images(text)
    text = _apply_backslash_escapes(text)
    # Быстрый путь: нет спецсимволов — один add_run без regex (с разбивкой по эмодзи)
    if '*' not in text and '_' not in text and '`' not in text and '~' not in text and '[' not in text:
        if 'http' in text or 'www.' in text:
            _split_and_add_with_bare_urls(paragraph, text, footnote_ctx)
        else:
            _add_run_with_emoji_font(paragraph, text)
        return
    # Защита: много * или _ — возможен тяжёлый backtracking, добавляем как текст
    if text.count('*') > 40 or text.count('_') > 40:
        _add_run_with_emoji_font(paragraph, text)
        return
    # Ограничиваем длину, чтобы избежать катастрофического backtracking
    _max_inline = 3000
    if len(text) > _max_inline:
        head, tail = text[: _max_inline], text[_max_inline:]
        _add_inline_formatted(paragraph, head, footnote_ctx)
        _add_run_with_emoji_font(paragraph, tail)
        return
    parts = _RE_INLINE.findall(text)
    if not parts:
        _add_run_with_emoji_font(paragraph, text)
        return
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            _add_run_with_emoji_font(paragraph, part[3:-3], bold=True, italic=True)
        elif part.startswith('___') and part.endswith('___'):
            _add_run_with_emoji_font(paragraph, part[3:-3], bold=True, italic=True)
        elif part.startswith('**') and part.endswith('**'):
            _add_run_with_emoji_font(paragraph, part[2:-2], bold=True)
        elif part.startswith('__') and part.endswith('__'):
            _add_run_with_emoji_font(paragraph, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            _add_run_with_emoji_font(paragraph, part[1:-1], italic=True)
        elif part.startswith('_') and part.endswith('_') and len(part) > 1:
            _add_run_with_emoji_font(paragraph, part[1:-1], italic=True)
        elif part.startswith('`') and part.endswith('`'):
            _add_run_with_emoji_font(paragraph, part[1:-1], font_name='Consolas')
        elif part.startswith('~~') and part.endswith('~~'):
            _add_run_with_emoji_font(paragraph, part[2:-2], strike=True)
        elif part.startswith('[^') and part.endswith(']'):
            m = _RE_FOOTNOTE_REF.match(part)
            if m and footnote_ctx:
                fn_id = m.group(1)
                entries = footnote_ctx['entries']
                defs = footnote_ctx['defs']
                if fn_id not in footnote_ctx['id_to_num']:
                    footnote_ctx['id_to_num'][fn_id] = len(entries) + 1
                    entries.append((fn_id, defs.get(fn_id, '')))
                num = footnote_ctx['id_to_num'][fn_id]
                run = paragraph.add_run(f' [{num}]')
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)
                run.font.superscript = True
            else:
                _add_run_with_emoji_font(paragraph, part)
        elif part.startswith('[') and '](' in part:
            m = _RE_LINK.match(part)
            if m:
                link_text = _strip_emoji(m.group(1))
                url = m.group(2)
                verified = (footnote_ctx or {}).get('verified_urls')
                company_url = (footnote_ctx or {}).get('company_url')
                if _url_is_verified(url, verified, company_url):
                    if link_text:
                        _add_hyperlink(paragraph, link_text, url)
                    else:
                        _add_hyperlink(paragraph, url, url)
                else:
                    _add_run_with_emoji_font(paragraph, link_text or url)
            else:
                _add_run_with_emoji_font(paragraph, part)
        elif part.startswith(('http://', 'https://', 'www.')):
            url = _normalize_url(part.rstrip('.,;:!?'))
            verified = (footnote_ctx or {}).get('verified_urls')
            company_url = (footnote_ctx or {}).get('company_url')
            if _url_is_verified(url, verified, company_url):
                _add_hyperlink(paragraph, part, url)
            else:
                _add_run_with_emoji_font(paragraph, part)
        else:
            _split_and_add_with_bare_urls(paragraph, part, footnote_ctx)


def _normalize_task_list_text(text):
    """Если текст начинается с [ ], [x] или [X], заменяет на [x]/[ ] (эмодзи убраны)."""
    if not text:
        return text
    m = re.match(r'^\[([ xX])\]\s*', text.strip())
    if m:
        checked = m.group(1).lower() == 'x'
        rest = text.strip()[len(m.group(0)):]
        return ('[x] ' if checked else '[ ] ') + rest
    return text


def _get_line_type(stripped):
    """Определяет тип строки: heading, list или para. Для решения, вставлять ли пустую строку."""
    if not stripped:
        return None
    if re.match(r'^#{1,6}\s+', stripped):
        return 'heading'
    if re.match(r'^[-*+]\s+', stripped) or re.match(r'^\d+\.\s+', stripped):
        return 'list'
    if _line_starts_with_emoji_bullet(stripped)[0]:  # (is_emoji_bullet, bullet, rest)
        return 'list'
    return 'para'


def _parse_table_row(line):
    """Парсит строку Markdown-таблицы в список ячеек (текст без лишних пробелов)."""
    stripped = line.strip()
    if not stripped or '|' not in stripped:
        return []
    parts = stripped.split('|')
    # Убираем пустые крайние от разделителей: | a | b | -> ['', ' a ', ' b ', '']
    cells = [p.strip() for p in parts[1:-1]] if len(parts) >= 3 else []
    return cells


def _is_table_separator_row(cells):
    """Проверяет, является ли строка разделителем таблицы (|---|---|)."""
    if not cells:
        return False
    return all(re.match(r'^:?-+:?$', c.strip()) for c in cells)


def _parse_col_alignment_from_separator(cell):
    """Парсит выравнивание из ячейки разделителя: :--- = left, :---: = center, ---: = right."""
    c = cell.strip()
    if c.startswith(':') and c.endswith(':'):
        return WD_ALIGN_PARAGRAPH.CENTER
    if c.endswith(':'):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _table_row_like(line):
    """Проверяет, похожа ли строка на строку таблицы (начинается с | и есть ещё |)."""
    s = line.strip()
    return len(s) >= 2 and s.startswith('|') and s.count('|') >= 2


def _set_cell_shading(cell, fill_hex):
    """Заливка ячейки цветом (hex без #, например D6EBFF)."""
    if not DOCX_AVAILABLE:
        return
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), fill_hex)
        tcPr.append(shading)
    except Exception:
        pass


def _set_cell_no_wrap(cell):
    """Запрет переноса строки в ячейке."""
    if not DOCX_AVAILABLE:
        return
    try:
        tc = cell._tc
        if not tc.xpath(".//w:noWrap"):
            tcPr = tc.get_or_add_tcPr()
            noWrap = OxmlElement('w:noWrap')
            tcPr.append(noWrap)
    except Exception:
        pass


def _set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    """Отступы ячейки в dxa (1 pt = 20 dxa). 90 dxa ≈ 4.5 pt."""
    if not DOCX_AVAILABLE:
        return
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in list(tcPr.xpath(".//w:tcMar")):
            old.getparent().remove(old)
        tcMar = OxmlElement('w:tcMar')
        for name, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            node = OxmlElement(f'w:{name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
    except Exception:
        pass


def _set_table_header_repeat(table):
    """Повтор первой строки таблицы на каждой странице."""
    if not DOCX_AVAILABLE:
        return
    try:
        first_row = table.rows[0]._tr
        trPr = first_row.get_or_add_trPr()
        if not trPr.xpath(".//w:tblHeader"):
            trPr.append(OxmlElement('w:tblHeader'))
    except Exception:
        pass


def _md_to_docx_content(doc, md_text, spacing=None, options=None, verified_urls=None, company_url=None):
    """Заполняет документ python-docx контентом из Markdown. verified_urls — только эти URL делаются кликабельными. company_url — для проверки по домену если verified_urls нет."""
    t0 = time.perf_counter()
    def_pt = lambda d, key, subkey: (d or {}).get(key, {}).get(subkey, 0)
    if spacing is None:
        spacing = {}
    opts = options or {}
    line_spacing = opts.get('line_spacing', 1.15)
    table_font_size = opts.get('table_font_size', 9)
    main_font_size = opts.get('main_font_size', 11)
    alternating_rows = opts.get('alternating_rows', True)
    use_hyphen_marker = (opts.get('list_marker', 'default') == 'hyphen')
    main_pt = max(8, min(24, int(main_font_size) if main_font_size else 11))
    # Иерархия кегля: обычный текст меньше любого заголовка
    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(main_pt)
    doc.styles['Normal'].paragraph_format.space_before = Pt(def_pt(spacing, 'normal', 'before'))
    doc.styles['Normal'].paragraph_format.space_after = Pt(def_pt(spacing, 'normal', 'after'))
    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles['Normal'].paragraph_format.line_spacing = line_spacing
    for level, size_pt in [(1, 16), (2, 14), (3, 13), (4, 12)]:
        try:
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = 'Segoe UI'
            h_style.font.size = Pt(size_pt)
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h_style.paragraph_format.space_before = Pt(def_pt(spacing, f'heading{level}', 'before'))
            h_style.paragraph_format.space_after = Pt(def_pt(spacing, f'heading{level}', 'after'))
        except KeyError:
            pass
    # Страница A4, поля: слева 2 см, справа 1,5 см
    try:
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
    except Exception:
        pass
    lines = md_text.replace('\r\n', '\n').replace('\r', '\n').split('\n')
    # Собираем определения сносок [^id]: text и удаляем их из документа
    footnote_defs = {}
    lines_no_fndefs = []
    for line in lines:
        m = _RE_FOOTNOTE_DEF.match(line.strip())
        if m:
            footnote_defs[m.group(1)] = m.group(2).strip()
        else:
            lines_no_fndefs.append(line)
    lines = lines_no_fndefs
    footnote_ctx = {'defs': footnote_defs, 'entries': [], 'id_to_num': {}, 'verified_urls': verified_urls, 'company_url': company_url}
    n_lines = len(lines)
    i = 0
    in_fence = False
    fence_char = None
    code_lines = []
    last_type = None
    _log = None
    try:
        from logger import logger as _log
    except ImportError:
        pass
    iter_limit = n_lines + 200
    while i < len(lines):
        iter_limit -= 1
        if iter_limit <= 0:
            if _log:
                _log.error("md-to-docx: превышен лимит итераций, возможный бесконечный цикл")
            print("md-to-docx: ОШИБКА — превышен лимит итераций", flush=True)
            break
        if i % 10 == 0 or i == 1:
            msg = f"md-to-docx: строка {i+1}/{n_lines}, прошло {time.perf_counter() - t0:.2f} с"
            if _log:
                _log.info(msg)
            print(msg, flush=True)
        line = lines[i]
        stripped_line = line.strip()
        # Ограждённый блок кода: ``` или ~~~ (CommonMark)
        is_fence_start = stripped_line.startswith('```') or stripped_line.startswith('~~~')
        if is_fence_start:
            fence_prefix = stripped_line[:3]
            if not in_fence:
                in_fence = True
                fence_char = fence_prefix
                code_lines = []
            else:
                if stripped_line.startswith(fence_char):
                    if code_lines:
                        p = doc.add_paragraph()
                        p.style = 'Normal'
                        run = p.add_run('\n'.join(code_lines))
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    in_fence = False
            i += 1
            last_type = 'code'
            continue
        if in_fence:
            code_lines.append(line)
            i += 1
            continue
        stripped = line.strip()
        if not stripped:
            # Пустые строки в MD пропускаем — заголовки и блоки имеют свои отступы
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            i = j
            continue
        # Горизонтальная линия (thematic break): только один тип символа, 3+ шт. (CommonMark)
        if re.match(r'^(\*{3,}|_{3,}|-{3,})\s*$', stripped):
            p_hr = doc.add_paragraph()
            p_hr.paragraph_format.space_before = Pt(6)
            p_hr.paragraph_format.space_after = Pt(6)
            try:
                p_hr.paragraph_format.border_bottom.width = Pt(0.5)
                p_hr.paragraph_format.border_bottom.color = RGBColor(0xC0, 0xC0, 0xC0)
            except Exception:
                pass
            i += 1
            last_type = 'para'
            continue
        # Цитата (blockquote): одна или несколько строк, начинающихся с >
        if stripped.startswith('>'):
            quote_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('>'):
                q = lines[j].strip()
                q = re.sub(r'^>\s*', '', q)
                quote_lines.append(q)
                j += 1
            p_q = doc.add_paragraph()
            p_q.paragraph_format.left_indent = Pt(24)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(0)
            for idx, qline in enumerate(quote_lines):
                if idx > 0:
                    p_q.add_run().add_break()
                _add_inline_formatted(p_q, qline, footnote_ctx)
            for run in p_q.runs:
                run.italic = True
            i = j
            last_type = 'para'
            continue
        # Индентационный блок кода (4 пробела или таб)
        if (line.startswith('    ') or line.startswith('\t')) and not in_fence:
            code_lines_indent = []
            j = i
            while j < len(lines) and (lines[j].startswith('    ') or lines[j].startswith('\t')):
                ln = lines[j]
                code_lines_indent.append(ln[4:] if ln.startswith('    ') else (ln[1:] if ln.startswith('\t') else ln))
                j += 1
            if code_lines_indent:
                p_code = doc.add_paragraph()
                p_code.style = 'Normal'
                run_code = p_code.add_run('\n'.join(code_lines_indent))
                run_code.font.name = 'Consolas'
                run_code.font.size = Pt(10)
            i = j
            last_type = 'code'
            continue
        # Заголовки раздела I. II. III. (римские цифры) — Heading 2 с отступом от предыдущего абзаца
        rm = re.match(r'^(I|II|III|IV|V|VI|VII|VIII|IX|X)\.\s+(.+)$', stripped, re.IGNORECASE)
        if rm:
            head_text = rm.group(2).strip()
            if head_text:
                p_h = doc.add_heading(head_text, level=2)
                p_h.clear()
                _add_inline_formatted(p_h, head_text, footnote_ctx)
                for run in p_h.runs:
                    run.font.size = Pt(14)
                pf = p_h.paragraph_format
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)
                i += 1
                last_type = 'heading'
                continue
        # Заголовки # ## ### — с поддержкой bold, ссылок и т.д.
        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            level = min(len(m.group(1)), 4)
            head_text = m.group(2).strip()
            p_h = doc.add_heading(head_text, level=level)
            p_h.clear()
            _add_inline_formatted(p_h, head_text, footnote_ctx)
            for run in p_h.runs:
                run.font.size = Pt([16, 14, 13, 12][min(level, 4) - 1])
            pf = p_h.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            i += 1
            last_type = 'heading'
            continue
        # Нумерованный список
        if re.match(r'^\d+\.\s+', stripped):
            text = _normalize_task_list_text(_strip_leading_bullet_chars(re.sub(r'^\d+\.\s+', '', stripped)))
            p = doc.add_paragraph(style='List Number')
            pf_list = p.paragraph_format
            pf_list.space_before = Pt(0)
            pf_list.space_after = Pt(3)
            _add_inline_formatted(p, text, footnote_ctx)
            # Пробуем слить следующую строку как расшифровку в тот же пункт (через перенос строки)
            next_i = i + 1
            if next_i < len(lines):
                next_raw = lines[next_i]
                next_stripped = next_raw.strip()
                if next_stripped and _get_line_type(next_stripped) == 'para':
                    run_br = p.add_run()
                    run_br.add_break()
                    _add_inline_formatted(p, next_stripped, footnote_ctx)
                    i = next_i + 1
                    last_type = 'list'
                    continue
            i += 1
            last_type = 'list'
            continue
        # Маркированный список -, * или + (CommonMark). Строки, оканчивающиеся на : — заголовки, не список (кроме подсписков и соседних пунктов)
        if re.match(r'^[-*+]\s+', stripped) or re.match(r'^\s{0,3}[-*+]\s+', line):
            text = _normalize_task_list_text(_strip_leading_bullet_chars(re.sub(r'^\s*[-*+]\s+', '', stripped)))
            curr_indent = len(line) - len(line.lstrip())
            next_line = lines[i + 1] if i + 1 < len(lines) else ''
            next_indent = len(next_line) - len(next_line.lstrip()) if next_line else -1
            next_is_sublist = (i + 1 < len(lines) and next_indent > curr_indent and re.match(r'^\s+[-*+]\s+', next_line))
            next_is_sibling_list = (i + 1 < len(lines) and next_indent == curr_indent and _get_line_type((next_line or '').strip()) == 'list')
            if text.endswith(':') and not next_is_sublist and not next_is_sibling_list:
                p = doc.add_paragraph()
                _add_inline_formatted(p, text, footnote_ctx)
                i += 1
                last_type = 'para'
                continue
            p = doc.add_paragraph(style='List Bullet' if not use_hyphen_marker else None)
            if use_hyphen_marker:
                run_marker = p.add_run('- ')
                run_marker.font.name = 'Segoe UI'
                run_marker.font.size = Pt(11)
            pf_list = p.paragraph_format
            pf_list.space_before = Pt(0)
            pf_list.space_after = Pt(3)
            _add_inline_formatted(p, text, footnote_ctx)
            # Пробуем слить следующую строку как расшифровку в тот же пункт (через перенос строки)
            next_i = i + 1
            if next_i < len(lines):
                next_raw = lines[next_i]
                next_stripped = next_raw.strip()
                if next_stripped and _get_line_type(next_stripped) == 'para':
                    run_br = p.add_run()
                    run_br.add_break()
                    _add_inline_formatted(p, next_stripped, footnote_ctx)
                    i = next_i + 1
                    last_type = 'list'
                    continue
            i += 1
            last_type = 'list'
            continue
        # Буллеты-эмодзи/символы (✅ текст, • пункт и т.п.) — маркер по настройке
        is_emoji_bullet, bullet, rest = _line_starts_with_emoji_bullet(stripped)
        if is_emoji_bullet and rest is not None:
            rest_text = _normalize_task_list_text(_strip_leading_bullet_chars(rest))
            curr_indent_e = len(line) - len(line.lstrip())
            next_line_e = lines[i + 1] if i + 1 < len(lines) else ''
            next_indent_e = len(next_line_e) - len(next_line_e.lstrip()) if next_line_e else -1
            next_is_sublist = (i + 1 < len(lines) and next_indent_e > curr_indent_e and re.match(r'^\s+[-*+]\s+', next_line_e))
            next_is_sibling_list = (i + 1 < len(lines) and next_indent_e == curr_indent_e and _get_line_type((next_line_e or '').strip()) == 'list')
            if rest_text.endswith(':') and not next_is_sublist and not next_is_sibling_list:
                p = doc.add_paragraph()
                _add_inline_formatted(p, rest_text, footnote_ctx)
                i += 1
                last_type = 'para'
                continue
            p = doc.add_paragraph(style='List Bullet' if not use_hyphen_marker else None)
            if use_hyphen_marker:
                run_marker = p.add_run('- ')
                run_marker.font.name = 'Segoe UI'
                run_marker.font.size = Pt(11)
            pf_list = p.paragraph_format
            pf_list.space_before = Pt(0)
            pf_list.space_after = Pt(3)
            _add_inline_formatted(p, rest_text, footnote_ctx)
            # Пробуем слить следующую строку как расшифровку в тот же пункт (через перенос строки)
            next_i = i + 1
            if next_i < len(lines):
                next_raw = lines[next_i]
                next_stripped = next_raw.strip()
                if next_stripped and _get_line_type(next_stripped) == 'para':
                    run_br = p.add_run()
                    run_br.add_break()
                    _add_inline_formatted(p, next_stripped, footnote_ctx)
                    i = next_i + 1
                    last_type = 'list'
                    continue
            i += 1
            last_type = 'list'
            continue
        # Строка только с картинкой (часто под буллетом) — дописываем к предыдущему параграфу
        if re.fullmatch(r'\s*!\[[^\]]*\]\([^)]+\)\s*', stripped) and len(doc.paragraphs) > 0:
            last_p = doc.paragraphs[-1]
            last_p.add_run(' ')
            _add_inline_formatted(last_p, stripped, footnote_ctx)
            i += 1
            continue
        # Markdown-таблица: собираем все строки таблицы подряд
        if _table_row_like(line):
            table_rows_raw = []
            j = i
            while j < len(lines) and _table_row_like(lines[j]):
                table_rows_raw.append(lines[j])
                j += 1
            # Парсим строки в ячейки
            rows_cells = [_parse_table_row(r) for r in table_rows_raw]
            # Выравнивание из строки-разделителя (|:---|:---:|---:|)
            col_alignments = []
            for cells in rows_cells:
                if _is_table_separator_row(cells):
                    col_alignments = [_parse_col_alignment_from_separator(c) for c in cells]
                    break
            # Убираем строки-разделители (|---|---|)
            data_rows = [cells for cells in rows_cells if not _is_table_separator_row(cells)]
            if data_rows:
                num_rows = len(data_rows)
                num_cols = max(len(cells) for cells in data_rows) if data_rows else 0
                if num_cols > 0:
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    # Очищённый текст по ячейкам для расчёта ширины столбцов
                    cell_texts = []
                    for ri, cells in enumerate(data_rows):
                        row_texts = []
                        for ci, cell_text in enumerate(cells):
                            if ci < num_cols:
                                t = _replace_md_images(cell_text).strip()
                                t = re.sub(r'\*\*([^*]{1,500})\*\*', r'\1', t)
                                row_texts.append(t)
                            else:
                                row_texts.append('')
                        while len(row_texts) < num_cols:
                            row_texts.append('')
                        cell_texts.append(row_texts)
                    # Для расчёта ширины столбцов используем отображаемый текст (ссылки → только текст)
                    cell_texts_for_width = [[_text_for_width_calc(cell_texts[ri][ci]) for ci in range(num_cols)] for ri in range(num_rows)]
                    # Ширина столбца как в MD; приоритет — даты/год, «Образование слова», «Обоснование»
                    pt_per_char = 6.5
                    padding_per_col = 24
                    _date_in_re = re.compile(r'\d{4}-\d{2}-\d{2}|\d{1,2}\.\d{1,2}\.\d{4}|\d{1,2}/\d{1,2}/\d{4}')
                    _year_in_re = re.compile(r'(?<!\d)\d{4}(?!\d)|\d{4}\s+г\.?')
                    _sum_re = re.compile(r'^[\d\s.,]+$')
                    header_row_texts = [cell_texts[0][ci] if ci < len(cell_texts[0]) else '' for ci in range(num_cols)]
                    col_content_pt = []
                    for ci in range(num_cols):
                        max_len = max(len(cell_texts_for_width[ri][ci]) for ri in range(num_rows))
                        is_short_col = max_len <= 10
                        if is_short_col:
                            w = max(max_len * pt_per_char + 10, 55)
                        else:
                            w = max_len * pt_per_char + padding_per_col
                        cells = [cell_texts_for_width[ri][ci].strip() for ri in range(num_rows) if cell_texts_for_width[ri][ci].strip()]
                        has_date = any(_date_in_re.search(c) for c in cells)
                        has_year = any(_year_in_re.search(c) for c in cells)
                        has_sum = any(_sum_re.match(c) for c in cells)
                        hdr = (header_row_texts[ci] or '').lower()
                        has_obraz = 'образование' in hdr or 'слово' in hdr
                        has_obosn = 'обоснование' in hdr
                        if is_short_col:
                            base = w
                            weight = 1.0
                        else:
                            min_pt = 32
                            if has_date or has_year:
                                min_pt = max(min_pt, 55)
                            if has_sum:
                                min_pt = max(min_pt, 55)
                            base = max(w, min_pt)
                            weight = 1.5 if (has_date or has_year) else 1.0
                            if has_obraz or has_obosn:
                                weight = max(weight, 1.4)
                        col_content_pt.append((base * weight, is_short_col))
                    try:
                        table.autofit = False
                        try:
                            table.allow_autofit = False
                        except Exception:
                            pass
                        content_width_pt = (21 - 2 - 1.5) * (72 / 2.54)  # область контента A4 в pt
                        col_vals = [x[0] for x in col_content_pt]
                        short_cols = [x[1] for x in col_content_pt]
                        total_content_pt = sum(col_vals)
                        # Если не вмещается — распределяем, но короткие столбцы (≤10 символов) не сжимаем
                        if total_content_pt > content_width_pt and total_content_pt > 0:
                            reserved = sum(col_vals[ci] for ci in range(num_cols) if short_cols[ci])
                            remaining = content_width_pt - reserved
                            long_cols_total = sum(col_vals[ci] for ci in range(num_cols) if not short_cols[ci])
                            if remaining > 0 and long_cols_total > 0:
                                col_widths_pt = []
                                for ci in range(num_cols):
                                    if short_cols[ci]:
                                        col_widths_pt.append(col_vals[ci])
                                    else:
                                        col_widths_pt.append(remaining * (col_vals[ci] / long_cols_total))
                            else:
                                col_widths_pt = [content_width_pt * (w / total_content_pt) for w in col_vals]
                        else:
                            col_widths_pt = list(col_vals)
                        # Ширину задаём в дюймах (1 pt = 1/72 inch) и для каждой ячейки столбца — иначе Word выравнивает столбцы одинаково
                        for ci in range(num_cols):
                            w_inches = col_widths_pt[ci] / 72.0
                            table.columns[ci].width = Inches(w_inches)
                            for ri in range(num_rows):
                                table.rows[ri].cells[ci].width = Inches(w_inches)
                        table.width = Inches(sum(col_widths_pt) / 72.0)
                    except Exception:
                        pass
                    # Размер шрифта: 8 pt при 10+ столбцах, иначе 10 pt заголовок / table_font_size тело
                    font_header_pt = 8 if num_cols >= 10 else 10
                    font_body_pt = 8 if num_cols >= 10 else table_font_size
                    _set_table_header_repeat(table)
                    for ri, row_texts in enumerate(cell_texts):
                        for ci, t in enumerate(row_texts):
                            if ci >= num_cols:
                                continue
                            cell = table.rows[ri].cells[ci]
                            para = cell.paragraphs[0]
                            para.clear()
                            if ri == 0:
                                raw_hdr = _replace_md_images(data_rows[0][ci]).strip() if ci < len(data_rows[0]) else t
                                _add_inline_formatted(para, raw_hdr, footnote_ctx)
                            else:
                                _add_inline_formatted(para, t)
                            for run in para.runs:
                                run.font.size = Pt(font_header_pt if ri == 0 else font_body_pt)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            _set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
                            if ri == 0:
                                pass
                            elif alternating_rows and num_rows >= 15 and ri % 2 == 1:
                                _set_cell_shading(cell, 'F5F5F5')
                            display_len = len(_text_for_width_calc(t))
                            if short_cols[ci] or (display_len <= 10) or (ri > 0 and _sum_re.match(t.strip()) and len(t.strip()) < 25):
                                _set_cell_no_wrap(cell)
                            align = col_alignments[ci] if ci < len(col_alignments) else WD_ALIGN_PARAGRAPH.LEFT
                            para.alignment = align
                            try:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                            except Exception:
                                pass
                    last_type = 'para'
            i = j
            continue
        # Setext-заголовки (CommonMark): строка(и) текста + подчёркивание === (h1) или --- (h2)
        setext_content = [stripped]
        j = i + 1
        found_setext = False
        while j < len(lines) and lines[j].strip():
            n = lines[j].strip()
            if re.match(r'^=+\s*$', n):
                head_text = ' '.join(setext_content)
                p_h = doc.add_heading(head_text, level=2)
                p_h.clear()
                _add_inline_formatted(p_h, head_text, footnote_ctx)
                for run in p_h.runs:
                    run.font.size = Pt(14)
                pf = p_h.paragraph_format
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)
                i = j + 1
                last_type = 'heading'
                found_setext = True
                break
            if re.match(r'^\-+\s*$', n) and len(n) >= 3:
                head_text = ' '.join(setext_content)
                p_h = doc.add_heading(head_text, level=2)
                p_h.clear()
                _add_inline_formatted(p_h, head_text, footnote_ctx)
                for run in p_h.runs:
                    run.font.size = Pt(14)
                pf = p_h.paragraph_format
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)
                i = j + 1
                last_type = 'heading'
                found_setext = True
                break
            if _get_line_type(n) != 'para':
                break
            setext_content.append(n)
            j += 1
        if not found_setext:
            # Обычный параграф (нет setext подчёркивания или прервались на заголовке/списке)
            p = doc.add_paragraph()
            _add_inline_formatted(p, stripped, footnote_ctx)
            i += 1
            last_type = 'para'
        continue
    # Раздел «Примечания» со сносками (ручная нумерация 1. 2. 3. — не продолжение списка)
    if footnote_ctx['entries']:
        doc.add_paragraph()
        doc.add_heading('Примечания', level=2)
        for num, (fn_id, fn_text) in enumerate(footnote_ctx['entries'], 1):
            p_fn = doc.add_paragraph()
            p_fn.paragraph_format.space_before = Pt(0)
            p_fn.paragraph_format.space_after = Pt(3)
            run_num = p_fn.add_run(f'{num}. ')
            run_num.font.name = 'Segoe UI'
            run_num.font.size = Pt(11)
            _add_inline_formatted(p_fn, fn_text, None)
    elapsed = time.perf_counter() - t0
    msg = f"md-to-docx: разбор завершён, {n_lines} строк за {elapsed:.2f} с"
    if _log:
        _log.info(msg)
    print(msg, flush=True)
    return doc

# Импорт модулей логирования и отслеживания расходов
try:
    from logger import logger
except ImportError:
    import logging
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)

try:
    from cost_tracker import get_balance, calculate_analysis_cost
    COST_TRACKING_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Модуль отслеживания расходов не доступен: {e}")
    COST_TRACKING_AVAILABLE = False

# #region agent log
try:
    import json
    from pathlib import Path
    debug_log_path = Path(__file__).parent / '.cursor' / 'debug.log'
    with open(debug_log_path, 'a', encoding='utf-8') as f:
        f.write(json.dumps({"sessionId":"debug-session","runId":"init","hypothesisId":"E","location":"main.py:64","message":"Попытка импорта Agents_crew","data":{"timestamp":__import__('datetime').datetime.now().isoformat()}})+'\n')
except: pass
# #endregion
# Для блока if __name__ (reloader, PID и т.д.)
is_werkzeug_main = os.environ.get('WERKZEUG_RUN_MAIN') == 'true'

# CrewAI — отложенная загрузка (при первом анализе), чтобы приложение запускалось даже при ошибках
crew = None
CREW_AVAILABLE = None
_crew_load_error = None

print("ℹ CrewAI загружается при первом анализе (отложенная инициализация)")


def _load_crew():
    """Загружает CrewAI при первом использовании. Возвращает (crew, error_msg)."""
    global crew, CREW_AVAILABLE, _crew_load_error
    if CREW_AVAILABLE is True and crew is not None:
        return crew, None
    if _crew_load_error and CREW_AVAILABLE is False:
        return None, _crew_load_error
    try:
        from Agents_crew import crew as c
        crew = c
        if crew is None:
            CREW_AVAILABLE = False
            _crew_load_error = (
                "CrewAI не инициализирован. Выполните:\n"
                "  pip install --upgrade \"crewai[tools]>=0.80\" \"pydantic>=2.10\"\n"
                "  python check_crewai.py"
            )
            return None, _crew_load_error
        CREW_AVAILABLE = True
        _crew_load_error = None
        logger.info("CrewAI успешно загружен")
        return crew, None
    except ImportError as e:
        CREW_AVAILABLE = False
        _crew_load_error = (
            f"CrewAI не установлен: {e}\n"
            "Решение: pip install --upgrade \"crewai[tools]>=0.80\" \"pydantic>=2.10\""
        )
        logger.error(_crew_load_error)
        return None, _crew_load_error
    except Exception as e:
        CREW_AVAILABLE = False
        _crew_load_error = (
            f"Ошибка CrewAI: {e}\n"
            "Попробуйте: pip install --upgrade \"crewai[tools]>=0.80\" \"pydantic>=2.10\""
        )
        logger.error(_crew_load_error, exc_info=True)
        return None, _crew_load_error

app = Flask(__name__)
CORS(app)

# Конфигурация
FLASK_HOST = os.getenv('FLASK_HOST', '0.0.0.0')
FLASK_PORT = int(os.getenv('FLASK_PORT', 5000))
FLASK_DEBUG = os.getenv('FLASK_DEBUG', 'True').lower() == 'true'

# Скрытый режим администратора: пароль для доступа к анализу сайта.
# Единый источник пароля — только APP_ACCESS_PASSWORD из окружения.
ADMIN_KEY = (os.getenv('ADMIN_KEY') or '123654+').strip()  # Код для входа в режим смены пароля
ENV_ACCESS_PASSWORD = (os.getenv('APP_ACCESS_PASSWORD') or '').strip()


def get_stored_password():
    """Возвращает пароль для доступа к анализу сайта."""
    return ENV_ACCESS_PASSWORD


def set_stored_password(new_password):
    """Пароль не меняется через API: используется только APP_ACCESS_PASSWORD в окружении."""
    raise RuntimeError('Смена пароля через API отключена. Измените APP_ACCESS_PASSWORD и перезапустите приложение.')


# Хранилище результатов анализа (в продакшене использовать Redis или БД)
analysis_results = {}
analysis_status = {}
analysis_costs = {}  # Хранилище расходов по задачам

@app.route('/')
def index():
    logger.info("Главная страница открыта")
    return render_template('index.html')


@app.route('/api/auth/check', methods=['POST'])
def auth_check():
    """Проверка пароля: верный пароль разблокирует анализ; 123654+ — вход в режим смены пароля."""
    if not request.json:
        return jsonify({'error': 'Требуется JSON'}), 400
    if not ENV_ACCESS_PASSWORD:
        return jsonify({'error': 'APP_ACCESS_PASSWORD не задан в окружении'}), 503
    password = (request.json.get('password') or '').strip()
    if password == ADMIN_KEY:
        return jsonify({'ok': True, 'admin': True})
    stored = get_stored_password()
    if password == stored:
        return jsonify({'ok': True, 'admin': False})
    return jsonify({'ok': False})


@app.route('/api/auth/change-password', methods=['POST'])
def auth_change_password():
    """Смена пароля доступа через API отключена (единый источник — env)."""
    if not request.json:
        return jsonify({'error': 'Требуется JSON'}), 400
    admin_key = (request.json.get('admin_key') or '').strip()
    new_password = (request.json.get('new_password') or '').strip()
    if admin_key != ADMIN_KEY:
        return jsonify({'error': 'Неверный код администратора'}), 403
    if not new_password:
        return jsonify({'error': 'Новый пароль не может быть пустым'}), 400
    try:
        set_stored_password(new_password)
    except RuntimeError as e:
        return jsonify({'error': str(e)}), 409
    return jsonify({'ok': True})


@app.route('/api/analyze', methods=['POST'])
def analyze():
    if not request.json:
        return jsonify({'error': 'Требуется JSON в теле запроса'}), 400
    
    data = request.json
    company_url = data.get('url', '').strip()
    
    if not company_url:
        return jsonify({'error': 'URL не может быть пустым'}), 400
    
    # Валидация и нормализация URL
    if not company_url.startswith(('http://', 'https://')):
        company_url = 'https://' + company_url
    
    # Проверка корректности URL
    try:
        parsed = urlparse(company_url)
        if not parsed.netloc:
            return jsonify({'error': 'Некорректный URL'}), 400
    except Exception:
        return jsonify({'error': 'Некорректный URL'}), 400
    
    # Проверка доступности CrewAI до создания задачи
    _c, _err = _load_crew()
    if _c is None:
        hint = ' Выполните: pip install --upgrade "crewai[tools]>=0.80" "pydantic>=2.10"'
        return jsonify({
            'error': 'CrewAI не доступен.' + hint
        }), 503
    
    # Генерируем уникальный ID для задачи
    task_id = datetime.now().strftime('%Y%m%d%H%M%S%f')
    
    logger.info(f"Запуск анализа для URL: {company_url} (Task ID: {task_id})")
    
    # Получаем начальный баланс для отслеживания расходов
    initial_balance = None
    if COST_TRACKING_AVAILABLE:
        try:
            initial_balance = get_balance()
            if initial_balance is not None:
                logger.info(f"Начальный баланс: {initial_balance} руб.")
            else:
                logger.warning("Не удалось получить начальный баланс. Отслеживание стоимости будет недоступно.")
        except Exception as e:
            logger.error(f"Ошибка при получении начального баланса: {e}")
            initial_balance = None
    
    # Устанавливаем статус "в процессе"
    analysis_status[task_id] = {
        'status': 'processing',
        'progress': 0,
        'message': 'Анализ запущен...',
        'cost': None,
        'initial_balance': initial_balance
    }
    
    # #region agent log
    write_debug_log({
        "sessionId": "debug-session",
        "runId": "task-created",
        "hypothesisId": "A",
        "location": "main.py:103",
        "message": "Задача создана",
        "data": {
            "task_id": task_id,
            "status": analysis_status[task_id],
            "timestamp": datetime.now().isoformat()
        }
    })
    # #endregion
    
    logger.info(f"Задача {task_id} создана в analysis_status. Всего задач: {len(analysis_status)}")
    
    # Запускаем анализ в отдельном потоке
    thread = threading.Thread(
        target=run_analysis,
        args=(task_id, company_url, initial_balance)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({
        'task_id': task_id,
        'status': 'processing',
        'message': 'Анализ запущен'
    })

def run_analysis(task_id, company_url, initial_balance=None):
    """Выполняет анализ сайта компании"""
    try:
        logger.info(f"[{task_id}] Начало анализа для {company_url}")
        
        # Проверяем, что задача существует
        if task_id not in analysis_status:
            logger.error(f"[{task_id}] Задача не найдена в analysis_status при запуске run_analysis")
            # Создаем задачу, если её нет
            analysis_status[task_id] = {
                'status': 'processing',
                'progress': 0,
                'message': 'Инициализация анализа...',
                'cost': None,
                'initial_balance': initial_balance
            }
            logger.warning(f"[{task_id}] Задача создана в run_analysis (не должна была отсутствовать)")
        
        _crew, err = _load_crew()
        if _crew is None:
            msg = (err or "CrewAI не доступен.").split('\n')[0]
            logger.error(f"[{task_id}] {msg}")
            raise Exception(msg)
        
        # Проверяем доступность сайта с правильными заголовками
        try:
            if task_id in analysis_status:
                update_progress_safely(task_id, 5, 'Проверка доступности сайта...')
            else:
                logger.error(f"[{task_id}] Задача исчезла из analysis_status перед проверкой сайта")
            
            # Заголовки для имитации обычного браузера
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
            }
            
            response = requests.get(company_url, timeout=10, allow_redirects=True, headers=headers)
            
            # Обрабатываем разные коды ответа
            if response.status_code == 403:
                # 403 (Forbidden) - сайт может блокировать автоматические запросы, но CrewAI может справиться
                logger.warning(f"[{task_id}] Сайт вернул 403 (Forbidden). Это может быть защита от ботов. Продолжаем анализ - CrewAI может обойти защиту.")
                # Не блокируем анализ, только предупреждаем
            elif response.status_code == 404:
                raise Exception("Сайт не найден (404). Проверьте правильность URL.")
            elif response.status_code >= 500:
                raise Exception(f"Ошибка сервера (код ответа: {response.status_code}). Сайт временно недоступен.")
            elif response.status_code >= 400:
                # Другие 4xx ошибки - предупреждаем, но продолжаем
                logger.warning(f"[{task_id}] Сайт вернул код {response.status_code}. Продолжаем анализ.")
            
            logger.info(f"[{task_id}] Проверка доступности завершена, код ответа: {response.status_code}")
            
        except requests.exceptions.Timeout:
            raise Exception("Сайт не отвечает (таймаут). Проверьте правильность URL и доступность сайта.")
        except requests.exceptions.ConnectionError as e:
            error_msg = str(e)
            if "resolve" in error_msg.lower() or "Name or service not known" in error_msg:
                raise Exception("Не удалось найти сайт. Проверьте правильность URL.")
            else:
                raise Exception("Не удалось подключиться к сайту. Проверьте правильность URL и доступность сайта.")
        except requests.exceptions.RequestException as e:
            error_msg = str(e)
            if "404" in error_msg or "not found" in error_msg.lower():
                raise Exception("Сайт не найден. Проверьте правильность URL.")
            else:
                # Для других ошибок предупреждаем, но не блокируем
                logger.warning(f"[{task_id}] Ошибка при проверке доступности: {error_msg}. Продолжаем анализ.")
        
        # Извлекаем название компании из URL
        parsed_url = urlparse(company_url)
        domain = parsed_url.netloc
        # Убираем www. и извлекаем основное имя домена
        domain_parts = domain.replace('www.', '').split('.')
        company_name = domain_parts[0].capitalize() if domain_parts else 'Company'
        
        logger.info(f"[{task_id}] Название компании: {company_name}")
        
        # Обновляем статус
        update_progress_safely(task_id, 10, 'Подготовка к анализу...')
        logger.info(f"[{task_id}] Этап 1: Подготовка к анализу")
        
        # Подготавливаем входные данные
        inputs = {
            "company_url": company_url,
            "company_name": company_name,
        }
        
        # Запускаем поток для плавного обновления прогресса во время выполнения CrewAI
        import time
        progress_thread_running = threading.Event()
        progress_thread_running.set()
        start_time = time.time()
        start_progress = 15
        target_progress = 88  # Увеличено до 88% для более плавного финала
        estimated_duration = 120  # Ожидаемая длительность анализа в секундах (2 минуты)
        
        def update_progress_during_analysis():
            """Плавно обновляет прогресс во время выполнения CrewAI анализа"""
            messages = [
                'Инициализация анализа...',
                'Сбор информации с сайта...',
                'Обработка данных...',
                'Анализ структуры компании...',
                'Анализ продуктов и услуг...',
                'Анализ финансовых показателей...',
                'Анализ стратегических направлений...',
                'Формирование аналитического отчета...',
                'Структурирование данных...',
                'Финальная обработка...',
            ]
            message_index = 0
            
            last_progress = start_progress  # Отслеживаем последнее значение прогресса
            
            while progress_thread_running.is_set():
                if task_id not in analysis_status or analysis_status[task_id]['status'] != 'processing':
                    break
                
                elapsed_time = time.time() - start_time
                
                # Более равномерное линейное увеличение прогресса
                if elapsed_time < estimated_duration:
                    # Линейное увеличение для равномерности
                    progress_ratio = min(elapsed_time / estimated_duration, 1.0)
                    # Используем слегка ускоренную функцию для более естественного движения
                    # Комбинация линейной и квадратичной для плавности без замедления в начале
                    smooth_ratio = progress_ratio * (0.3 + 0.7 * progress_ratio)  # 30% линейная + 70% квадратичная
                    calculated_progress = int(start_progress + (target_progress - start_progress) * smooth_ratio)
                else:
                    # Если прошло больше времени, медленно приближаемся к target_progress
                    extra_time = elapsed_time - estimated_duration
                    # Медленное увеличение после истечения ожидаемого времени
                    additional_progress = min(int(extra_time / 10), target_progress - start_progress)
                    calculated_progress = min(start_progress + additional_progress, target_progress)
                
                # ВАЖНО: Прогресс может только увеличиваться, никогда не уменьшаться
                # Используем max, чтобы гарантировать монотонное увеличение
                current_progress = max(last_progress, calculated_progress)
                last_progress = current_progress  # Сохраняем для следующей итерации
                
                # Обновляем сообщение каждые ~12 секунд (10 сообщений за ~120 секунд)
                if elapsed_time > 0:
                    new_message_index = min(int(elapsed_time / (estimated_duration / len(messages))), len(messages) - 1)
                    if new_message_index != message_index:
                        message_index = new_message_index
                
                if task_id in analysis_status and analysis_status[task_id]['status'] == 'processing':
                    update_progress_safely(task_id, current_progress, messages[message_index])
                
                time.sleep(1)  # Обновляем каждую секунду для плавности
        
        progress_thread = threading.Thread(target=update_progress_during_analysis)
        progress_thread.daemon = True
        progress_thread.start()
        
        # Запускаем анализ
        update_progress_safely(task_id, 15, 'Запуск анализа...')
        logger.info(f"[{task_id}] Этап 2: Запуск CrewAI анализа")
        
        # #region agent log
        write_debug_log({
            "sessionId": "debug-session",
            "runId": "pre-kickoff",
            "hypothesisId": "B",
            "location": "main.py:142",
            "message": "Перед вызовом crew.kickoff()",
            "data": {
                "task_id": task_id,
                "timestamp": datetime.now().isoformat()
            }
        })
        # #endregion
        
        result = _crew.kickoff(inputs=inputs)
        result_str = str(result)

        # Реальные пути с сайта (HTML-обход) — whitelist для ссылок и снятие выдуманных [текст](url)
        verified_urls = set()
        try:
            if task_id in analysis_status:
                update_progress_safely(task_id, analysis_status[task_id].get('progress', 88), 'Проверка структуры ссылок сайта...')
        except Exception:
            pass
        try:
            allowed_keys, verified_from_crawl = _discover_internal_site_urls(company_url)
            if allowed_keys:
                result_str = _sanitize_markdown_links(result_str, allowed_keys)
                verified_urls = verified_from_crawl
                logger.info(
                    f"[{task_id}] Обход сайта: {len(allowed_keys)} уникальных URL; "
                    f"в отчёте оставлены только ссылки из реальной структуры"
                )
            else:
                logger.warning(f"[{task_id}] Обход сайта не дал ссылок — whitelist только из Задачи 1 (слабее)")
                task1_path = Path(__file__).parent / 'tasks' / 'task_1_scraped_data.md'
                if task1_path.exists():
                    task1_text = task1_path.read_text(encoding='utf-8', errors='replace')
                    verified_urls = _extract_urls_from_markdown(task1_text)
                    t1_keys = {_canonical_url_key(u) for u in verified_urls}
                    result_str = _sanitize_markdown_links(result_str, t1_keys)
                    logger.info(f"[{task_id}] Извлечено {len(verified_urls)} URL из Task 1 для фильтрации")
        except Exception as e:
            logger.warning(f"[{task_id}] Ошибка обхода/фильтра ссылок: {e}")
            try:
                task1_path = Path(__file__).parent / 'tasks' / 'task_1_scraped_data.md'
                if task1_path.exists():
                    task1_text = task1_path.read_text(encoding='utf-8', errors='replace')
                    verified_urls = _extract_urls_from_markdown(task1_text)
            except Exception:
                pass
        
        # Останавливаем поток обновления прогресса
        progress_thread_running.clear()
        
        # #region agent log
        write_debug_log({
            "sessionId": "debug-session",
            "runId": "post-kickoff",
            "hypothesisId": "B",
            "location": "main.py:145",
            "message": "После вызова crew.kickoff()",
            "data": {
                "task_id": task_id,
                "timestamp": datetime.now().isoformat()
            }
        })
        # #endregion
        
        logger.info(f"[{task_id}] CrewAI анализ завершен")
        
        # Плавное завершение с небольшими шагами
        update_progress_safely(task_id, 90, 'Завершение анализа...')
        logger.info(f"[{task_id}] Этап 3: Завершение анализа")
        
        import time
        time.sleep(0.5)  # Небольшая пауза для плавности
        
        # Вычисляем стоимость анализа
        update_progress_safely(task_id, 93, 'Расчет стоимости...')
        cost = None
        if COST_TRACKING_AVAILABLE and initial_balance is not None:
            cost = calculate_analysis_cost(task_id, initial_balance)
            analysis_costs[task_id] = cost
            logger.info(f"[{task_id}] Стоимость анализа: {cost} руб.")
        
        time.sleep(0.3)  # Небольшая пауза для плавности
        
        # Сохраняем результат
        update_progress_safely(task_id, 96, 'Сохранение результатов...')
        analysis_results[task_id] = {
            'result': result_str,
            'url': company_url,
            'company_name': company_name,
            'timestamp': datetime.now().isoformat(),
            'cost': cost,
            'task_id': task_id,
            'verified_urls': list(verified_urls) if verified_urls else None
        }
        
        # Устанавливаем статус "завершено"
        analysis_status[task_id] = {
            'status': 'completed',
            'progress': 100,
            'message': 'Анализ завершен',
            'cost': cost,
            'initial_balance': initial_balance
        }
        
        logger.info(f"[{task_id}] Анализ успешно завершен. Готов к новой итерации.")
        
    except Exception as e:
        error_message = str(e)
        logger.error(f"[{task_id}] Ошибка при анализе: {error_message}", exc_info=True)
        analysis_status[task_id] = {
            'status': 'error',
            'progress': 0,
            'message': f'Ошибка: {error_message}',
            'cost': None
        }

@app.route('/api/status/<task_id>', methods=['GET'])
def get_status(task_id):
    # #region agent log
    write_debug_log({
        "sessionId": "debug-session",
        "runId": "status-check",
        "hypothesisId": "A",
        "location": "main.py:279",
        "message": "Запрос статуса задачи",
        "data": {
            "task_id": task_id,
            "exists": task_id in analysis_status,
            "all_tasks": list(analysis_status.keys())[-5:],
            "timestamp": datetime.now().isoformat()
        }
    })
    # #endregion
    
    if task_id not in analysis_status:
        logger.warning(f"Запрос статуса для несуществующей задачи: {task_id}. Доступные задачи: {list(analysis_status.keys())[-5:]}")
        return jsonify({'error': 'Задача не найдена'}), 404
    
    status = analysis_status[task_id].copy()
    
    # Если анализ завершен, добавляем результат
    if status['status'] == 'completed':
        if task_id in analysis_results:
            result_data = analysis_results[task_id].copy()
            status['result'] = result_data
            # Добавляем информацию о стоимости в статус и результат
            if 'cost' in result_data:
                cost_value = result_data['cost']
                status['cost'] = cost_value
                # Убеждаемся, что стоимость есть в результате
                if 'cost' not in result_data or result_data['cost'] is None:
                    result_data['cost'] = cost_value
            else:
                # Если стоимости нет в результате, но есть в статусе
                if 'cost' in status and status['cost'] is not None:
                    result_data['cost'] = status['cost']
            
            logger.info(f"[{task_id}] Возврат статуса: cost={status.get('cost')}, result.cost={result_data.get('cost')}")
    
    return jsonify(status)

@app.route('/api/export/<task_id>', methods=['GET'])
def export_to_docx(task_id):
    """Экспорт результатов анализа в DOCX формат"""
    if task_id not in analysis_results:
        return jsonify({'error': 'Результаты анализа не найдены'}), 404
    
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'Модуль python-docx не установлен'}), 500
    
    try:
        result_data = analysis_results[task_id]
        
        # Создаем документ
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Заголовок
        title = doc.add_heading('КОМПЛЕКСНЫЙ КОРПОРАТИВНЫЙ ОТЧЕТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о компании
        company_name = result_data.get('company_name', 'Неизвестная компания')
        company_url = result_data.get('url', 'N/A')
        timestamp = result_data.get('timestamp', datetime.now().isoformat())
        cost = result_data.get('cost')
        
        doc.add_paragraph(f'Компания: {company_name}')
        doc.add_paragraph(f'URL: {company_url}')
        doc.add_paragraph(f'Дата анализа: {timestamp}')
        if cost is not None:
            doc.add_paragraph(f'Стоимость анализа: {cost:.4f} руб.')
        
        # Пустая строка без дополнительных отступов
        p_blank_header = doc.add_paragraph()
        pf_header = p_blank_header.paragraph_format
        pf_header.space_before = Pt(0)
        pf_header.space_after = Pt(0)
        
        # Основной контент — полный разбор Markdown; только проверенные ссылки — кликабельные
        content = result_data.get('result', '')
        v_urls = result_data.get('verified_urls')
        verified_urls = set(v_urls) if v_urls else None
        company_url = result_data.get('url', '')
        if content:
            _md_to_docx_content(doc, content, spacing=None, options={'line_spacing': 1.15, 'main_font_size': 11},
                                verified_urls=verified_urls, company_url=company_url)
        
        # Сохраняем в память
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Формируем имя файла (ASCII для HTTP-заголовков)
        filename = f'analysis_{company_name}_{datetime.now().strftime("%Y%m%d")}.docx'
        safe_filename = "".join(c if ord(c) < 128 else '_' for c in filename).rstrip('_') or 'analysis.docx'
        
        logger.info(f"Экспорт результатов {task_id} в DOCX: {safe_filename}")
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=safe_filename
        )
        
    except Exception as e:
        logger.error(f"Ошибка при экспорте в DOCX для задачи {task_id}: {e}", exc_info=True)
        return jsonify({'error': f'Ошибка при создании документа: {str(e)}'}), 500


def _postprocess_pandoc_docx(docx_bytes, spacing=None, options=None):
    """Постобработка DOCX после Pandoc: отступы, межстрочный интервал, шрифт таблиц, чередование строк."""
    if not DOCX_AVAILABLE or not docx_bytes:
        return docx_bytes
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
    except Exception:
        return docx_bytes
    spacing = spacing or {}
    opts = options or {}
    line_spacing = opts.get('line_spacing', 1.15)
    table_font_size = opts.get('table_font_size', 9)
    main_font_size = opts.get('main_font_size', 11)
    alternating_rows = opts.get('alternating_rows', True)
    def_pt = lambda key, subkey: (spacing.get(key, {}) or {}).get(subkey, 0)
    main_pt = max(8, min(24, int(main_font_size) if main_font_size else 11))
    # Normal: отступы, межстрочный интервал, размер шрифта
    try:
        s = doc.styles['Normal']
        s.font.size = Pt(main_pt)
        s.paragraph_format.space_before = Pt(def_pt('normal', 'before'))
        s.paragraph_format.space_after = Pt(def_pt('normal', 'after'))
        s.paragraph_format.line_spacing = line_spacing
    except KeyError:
        pass
    # Heading 1-4: по левому краю
    for level, def_before, def_after in [(1, 12, 6), (2, 10, 4), (3, 8, 4), (4, 6, 2)]:
        try:
            s = doc.styles[f'Heading {level}']
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            s.paragraph_format.space_before = Pt(def_pt(f'heading{level}', 'before') or def_before)
            s.paragraph_format.space_after = Pt(def_pt(f'heading{level}', 'after') or def_after)
        except KeyError:
            pass
    # Таблицы: шрифт и чередование строк
    font_pt = max(6, min(24, int(table_font_size) if table_font_size else 9))
    for table in doc.tables:
        num_rows = len(table.rows)
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(font_pt)
            if alternating_rows and num_rows >= 15 and ri > 0 and ri % 2 == 1:
                for cell in row.cells:
                    _set_cell_shading(cell, 'F2F2F2')
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


_PANDOC_INPUT_FORMATS = {
    '.md': 'markdown', '.markdown': 'markdown',
    '.html': 'html', '.htm': 'html',
    '.rst': 'rst', '.rest': 'rst',
    '.tex': 'latex', '.latex': 'latex',
    '.txt': 'markdown', '.text': 'markdown',
}


def _convert_md_to_docx_pandoc(md_bytes, base_name, input_format='markdown'):
    """Конвертация в DOCX через Pandoc (MD, HTML, RST, TXT и др.)."""
    try:
        result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True, timeout=5)
        if result.returncode != 0:
            return None, 'Pandoc не найден. Установите: https://pandoc.org/installing.html'
    except FileNotFoundError:
        return None, 'Pandoc не установлен. Установите: https://pandoc.org/installing.html'
    except subprocess.TimeoutExpired:
        return None, 'Ошибка проверки Pandoc'
    with tempfile.TemporaryDirectory() as tmpdir:
        ext_map = {'markdown': '.md', 'html': '.html', 'rst': '.rst', 'latex': '.tex'}
        ext = ext_map.get(input_format, '.md')
        inp_path = Path(tmpdir) / ('input' + ext)
        docx_path = Path(tmpdir) / 'output.docx'
        inp_path.write_bytes(md_bytes)
        try:
            result = subprocess.run(
                ['pandoc', '-f', input_format, '-t', 'docx', '-o', str(docx_path), str(inp_path)],
                capture_output=True,
                text=True,
                timeout=MD_DOCX_TIMEOUT,
                cwd=tmpdir
            )
            if result.returncode != 0:
                err = (result.stderr or result.stdout or 'Неизвестная ошибка').strip()
                return None, f'Ошибка Pandoc: {err}'
            if not docx_path.exists():
                return None, 'Pandoc не создал файл'
            return docx_path.read_bytes(), None
        except subprocess.TimeoutExpired:
            return None, f'Конвертация заняла больше {MD_DOCX_TIMEOUT} с'


@app.route('/api/convert/md-to-docx-pandoc', methods=['POST'])
def convert_md_to_docx_pandoc():
    """Конвертация MD → DOCX через Pandoc (стандартный алгоритм)."""
    logger.info("POST /api/convert/md-to-docx-pandoc: запрос получен")
    uploaded = request.files.get('file')
    if not uploaded or uploaded.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400
    try:
        md_bytes = uploaded.read()
        logger.info(f"POST /api/convert/md-to-docx-pandoc: файл прочитан, {len(md_bytes)} байт")
    except Exception as e:
        logger.error(f"Ошибка чтения файла MD → DOCX (Pandoc): {e}", exc_info=True)
        return jsonify({'error': f'Не удалось прочитать файл: {str(e)}'}), 500
    try:
        md_text = md_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            md_text = md_bytes.decode('utf-8-sig')
        except UnicodeDecodeError:
            try:
                md_text = md_bytes.decode('cp1251')
            except (UnicodeDecodeError, LookupError):
                return jsonify({'error': 'Неподдерживаемая кодировка. Файл должен быть текстовым (UTF-8).'}), 400
        md_bytes = md_text.encode('utf-8')
    suf = Path(uploaded.filename).suffix.lower()
    in_fmt = _PANDOC_INPUT_FORMATS.get(suf, 'markdown')
    docx_bytes, err = _convert_md_to_docx_pandoc(md_bytes, Path(uploaded.filename).stem, in_fmt)
    if err:
        logger.warning(f"Pandoc конвертация: {err}")
        return jsonify({'error': err}), 500
    use_pandoc_default = request.form.get('use_pandoc_default') in ('1', 'true', 'yes')
    if not use_pandoc_default:
        spacing = None
        options = None
        try:
            raw = request.form.get('spacing')
            if raw:
                spacing = json.loads(raw)
        except (TypeError, ValueError):
            pass
        try:
            raw_opts = request.form.get('options')
            if raw_opts:
                options = json.loads(raw_opts)
        except (TypeError, ValueError):
            pass
        docx_bytes = _postprocess_pandoc_docx(docx_bytes, spacing=spacing, options=options)
    safe_name = "".join(c for c in Path(uploaded.filename).stem if c.isalnum() or c in (' ', '-', '_')).rstrip() or 'document'
    download_name = f"{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    download_name_ascii = "".join(c if ord(c) < 128 else '_' for c in download_name).rstrip('_') or 'document.docx'
    disp_ascii = f'attachment; filename="{download_name_ascii}"'
    try:
        disp_value = f'{disp_ascii}; filename*=UTF-8\'\'{quote(download_name, safe="")}'
    except Exception:
        disp_value = disp_ascii
    resp = Response(docx_bytes, status=200, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp.headers['Content-Disposition'] = disp_value
    resp.headers['Content-Length'] = str(len(docx_bytes))
    logger.info(f"POST /api/convert/md-to-docx-pandoc: готово, {len(docx_bytes)} байт")
    return resp


@app.route('/api/convert/md-to-docx', methods=['POST'])
def convert_md_to_docx():
    """Принимает MD-файл, конвертирует в DOCX и возвращает файл для скачивания."""
    logger.info("POST /api/convert/md-to-docx: запрос получен")
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'Модуль python-docx не установлен'}), 500
    uploaded = request.files.get('file')
    if not uploaded or uploaded.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400
    try:
        md_bytes = uploaded.read()
        logger.info(f"POST /api/convert/md-to-docx: файл прочитан, {len(md_bytes)} байт")
    except Exception as e:
        logger.error(f"Ошибка чтения файла MD → DOCX: {e}", exc_info=True)
        return jsonify({'error': f'Не удалось прочитать файл: {str(e)}'}), 500
    try:
        md_text = md_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            md_text = md_bytes.decode('utf-8-sig')
        except UnicodeDecodeError:
            try:
                md_text = md_bytes.decode('cp1251')
            except (UnicodeDecodeError, LookupError):
                return jsonify({
                    'error': 'Неподдерживаемая кодировка файла. Сохраните файл в UTF-8 и попробуйте снова.'
                }), 400
    spacing = None
    options = None
    try:
        raw = request.form.get('spacing')
        if raw:
            spacing = json.loads(raw)
    except (TypeError, ValueError):
        pass
    try:
        raw_opts = request.form.get('options')
        if raw_opts:
            options = json.loads(raw_opts)
    except (TypeError, ValueError):
        pass

    def do_convert():
        doc = Document()
        t1 = time.perf_counter()
        _md_to_docx_content(doc, md_text, spacing=spacing, options=options)
        t2 = time.perf_counter()
        file_stream = io.BytesIO()
        doc.save(file_stream)
        t3 = time.perf_counter()
        file_stream.seek(0)
        logger.info(f"md-to-docx: разбор {t2-t1:.2f} с, сохранение DOCX {t3-t2:.2f} с, всего {t3-t1:.2f} с")
        return file_stream.getvalue()

    try:
        logger.info("POST /api/convert/md-to-docx: начало конвертации в DOCX")
        if os.environ.get('MD_DOCX_SYNC') == '1':
            # Синхронный режим: прогресс виден в терминале (print), для отладки
            docx_bytes = do_convert()
        else:
            future = _executor_md_docx.submit(do_convert)
            docx_bytes = future.result(timeout=MD_DOCX_TIMEOUT)
        logger.info(f"POST /api/convert/md-to-docx: готово, отправка {len(docx_bytes)} байт")
        base_name = Path(uploaded.filename).stem
        safe_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_')).rstrip() or 'document'
        download_name = f"{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        # ASCII fallback — HTTP-заголовки latin-1 не поддерживают кириллицу
        download_name_ascii = "".join(c if ord(c) < 128 else '_' for c in download_name).rstrip('_') or 'document.docx'
        # RFC 5987: filename* для Unicode, filename — ASCII fallback
        disp_ascii = f'attachment; filename="{download_name_ascii}"'
        try:
            encoded = quote(download_name, safe='')
            disp_value = f'{disp_ascii}; filename*=UTF-8\'\'{encoded}'
        except Exception:
            disp_value = disp_ascii
        resp = Response(docx_bytes, status=200, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp.headers['Content-Disposition'] = disp_value
        resp.headers['Content-Length'] = str(len(docx_bytes))
        return resp
    except FuturesTimeoutError:
        logger.error("POST /api/convert/md-to-docx: таймаут конвертации")
        return jsonify({'error': f'Конвертация заняла больше {MD_DOCX_TIMEOUT} с. Упростите файл или разбейте его.'}), 503
    except Exception as e:
        logger.error(f"Ошибка конвертации MD → DOCX: {e}", exc_info=True)
        return jsonify({'error': f'Ошибка конвертации: {str(e)}. Попробуйте другой файл или упростите разметку.'}), 500


if __name__ == '__main__':
    # Проверяем, что это дочерний процесс Flask reloader (рабочий сервер)
    # WERKZEUG_RUN_MAIN устанавливается Flask только в дочернем процессе
    # Используем уже определенную переменную is_werkzeug_main из начала файла
    
    # Проверки PID файла и порта выполняем только в основном процессе
    # (до запуска reloader) или если debug отключен
    if not is_werkzeug_main or not FLASK_DEBUG:
        # Регистрируем обработчик для удаления PID файла при завершении
        atexit.register(remove_pid_file)
        
        # Обработка сигналов для корректного завершения
        def signal_handler(signum, frame):
            print_red("\n" + "=" * 60)
            print_red("Завершение работы приложения...")
            print_red("=" * 60)
            print("(Сообщение «Exception in thread» при остановке — нормально, можно игнорировать)")
            remove_pid_file()
            sys.exit(0)
        
        signal.signal(signal.SIGINT, signal_handler)
        signal.signal(signal.SIGTERM, signal_handler)
        
        # ШАГ 1: Проверяем наличие запущенного экземпляра приложения через PID файл
        # Это основная проверка - если приложение запущено, оно создало PID файл
        logger.info("Проверка запущенных экземпляров приложения...")
        print("Проверка запущенных экземпляров приложения...")
        if not check_and_kill_existing_instance():
            logger.error("Не удалось завершить предыдущий экземпляр приложения")
            print_red("❌ Ошибка: не удалось завершить предыдущий экземпляр приложения")
            sys.exit(1)
        
        # ========================================================================
        # ШАГ 2: Дополнительная проверка - занят ли порт
        # ========================================================================
        # В контейнере Docker пропускаем проверку порта - Docker сам управляет портами
        # В debug режиме Flask reloader создаст дочерний процесс, который займет порт - это нормально
        # Поэтому в debug режиме проверяем порт только один раз, перед запуском Flask
        if not is_running_in_docker() and not FLASK_DEBUG:
            # Проверяем порт только если debug отключен (нет Flask reloader)
            # Проверяем порт на случай, если PID файл был удален, но процесс еще работает
            # или если другой процесс использует порт 5000
            check_host = FLASK_HOST if FLASK_HOST != '0.0.0.0' else '127.0.0.1'
            import time
            
            # Проверяем порт несколько раз с попытками завершения
            # Увеличено количество попыток для более надежного освобождения порта
            max_attempts = 10
            for attempt in range(max_attempts):
                if check_port_in_use(check_host, FLASK_PORT):
                    if attempt == 0:
                        logger.warning(f"Порт {FLASK_PORT} уже занят. Принудительно завершаю предыдущий процесс...")
                        print_yellow(f"⚠ Порт {FLASK_PORT} уже занят. Принудительно завершаю предыдущий процесс...")
                    else:
                        logger.warning(f"Порт {FLASK_PORT} все еще занят (попытка {attempt + 1}/{max_attempts}). Повторная попытка завершения...")
                        print_yellow(f"⚠ Порт {FLASK_PORT} все еще занят. Повторная попытка завершения (попытка {attempt + 1}/{max_attempts})...")
                    
                    # Исключаем текущий процесс из проверки
                    kill_process_on_port(FLASK_PORT, exclude_pid=os.getpid())
                    
                    # Ждем, чтобы порт освободился (увеличено время ожидания для надежности)
                    time.sleep(3)  # Увеличено с 2 до 3 секунд
                    
                    # Проверяем снова
                    if not check_port_in_use(check_host, FLASK_PORT):
                        logger.info(f"Порт {FLASK_PORT} успешно освобожден после {attempt + 1} попыток")
                        print_green(f"✓ Порт {FLASK_PORT} успешно освобожден после {attempt + 1} попыток")
                        break
                else:
                    # Порт свободен, можно запускать
                    break
            else:
                # Все попытки исчерпаны, порт все еще занят
                logger.error(f"Порт {FLASK_PORT} все еще занят после {max_attempts} попыток завершения процесса")
                print_red(f"❌ Ошибка: порт {FLASK_PORT} все еще занят после {max_attempts} попыток.")
                print_red(f"   Закройте приложение вручную или измените порт в переменных окружения.")
                print_red(f"   Используйте скрипт scripts/stop_app.py для принудительного завершения.")
                sys.exit(1)
        elif not is_running_in_docker() and FLASK_DEBUG:
            # В debug режиме проверяем порт только один раз, перед запуском Flask
            # Flask reloader создаст дочерний процесс, который займет порт - это нормально
            # Проверяем только, не занят ли порт другим процессом (не нашим)
            check_host = FLASK_HOST if FLASK_HOST != '0.0.0.0' else '127.0.0.1'
            if check_port_in_use(check_host, FLASK_PORT):
                # Проверяем, не занят ли порт другим процессом (не нашим)
                # Исключаем текущий процесс из проверки
                killed = kill_process_on_port(FLASK_PORT, exclude_pid=os.getpid())
                if killed:
                    import time
                    time.sleep(1)  # Небольшая пауза для освобождения порта
                    logger.info("Освобожден порт от предыдущего процесса")
                    print_green("✓ Освобожден порт от предыдущего процесса")
                else:
                    # Порт занят, но это может быть наш дочерний процесс от предыдущего запуска
                    # Или другой процесс - проверяем через PID файл (уже сделано в ШАГ 1)
                    logger.info("Порт занят, но это может быть нормально в debug режиме (Flask reloader)")
                    print("ℹ Порт занят - в debug режиме Flask reloader создаст дочерний процесс")
        else:
            logger.info("Запуск в контейнере Docker - пропускаем проверку порта (Docker управляет портами)")
            print("ℹ Запуск в контейнере Docker - проверка порта не требуется")
        
        # ШАГ 3: Создаем PID файл для текущего экземпляра ПЕРЕД запуском
        # Это важно - флаг должен быть установлен до того, как приложение начнет слушать порт
        create_pid_file()
        logger.info(f"PID файл создан: {PID_FILE} (PID: {os.getpid()})")
    
    # В дочернем процессе также регистрируем обработчик для удаления PID файла
    if is_werkzeug_main:
        atexit.register(remove_pid_file)
    
    logger.info("=" * 60)
    logger.info("Запуск Flask приложения...")
    logger.info(f"Сервер будет доступен по адресу: http://{FLASK_HOST}:{FLASK_PORT}")
    logger.info("Нажмите Ctrl+C для остановки")
    logger.info("=" * 60)
    
    print_green("=" * 60)
    print_green("Запуск Flask приложения...")
    print_green(f"Сервер будет доступен по адресу: http://{FLASK_HOST}:{FLASK_PORT}")
    print_green("Нажмите Ctrl+C для остановки")
    print_green("=" * 60)
    
    # Обработчик Ctrl+C: KeyboardInterrupt — стандартный выход, atexit удалит PID
    def _shutdown_handler(signum, frame):
        print_red("\nЗавершение работы приложения...")
        remove_pid_file()
        raise KeyboardInterrupt()
    try:
        signal.signal(signal.SIGINT, _shutdown_handler)
        signal.signal(signal.SIGTERM, _shutdown_handler)
    except (ValueError, OSError):
        pass
    
    # На Windows с reloader Ctrl+C часто не доходит до рабочего процесса — отключаем reloader по умолчанию
    use_reloader = FLASK_DEBUG and (platform.system() != 'Windows' or os.environ.get('FLASK_USE_RELOADER') == '1')
    app.run(debug=FLASK_DEBUG, host=FLASK_HOST, port=FLASK_PORT, use_reloader=use_reloader, threaded=True)

